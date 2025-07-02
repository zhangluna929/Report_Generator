import os
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.pdfgen import canvas

# 数据加载与预处理
def load_experiment_data(file_path):
    """
    加载并返回实验数据。
    :param file_path: 数据文件路径
    :return: 返回 pandas DataFrame 格式的数据
    """
    try:
        data = pd.read_csv(file_path)
    except Exception as e:
        print(f"加载数据时出错: {e}")
        raise
    return data

def generate_plots(data, output_dir):
    """
    生成实验数据的图表。
    :param data: 实验数据
    :param output_dir: 图表输出文件夹路径
    :return: 返回图表的路径
    """
    # 电压与时间的关系图
    voltage_plot_path = os.path.join(output_dir, 'voltage_vs_time.png')
    plt.figure(figsize=(8, 6))
    plt.plot(data['time'], data['voltage'], label='Voltage', color='b')
    plt.xlabel('Time (s)')
    plt.ylabel('Voltage (V)')
    plt.title('Voltage vs Time')
    plt.grid(True)
    plt.legend()
    plt.savefig(voltage_plot_path)
    plt.close()

    # 温度与时间的关系图
    temperature_plot_path = os.path.join(output_dir, 'temperature_vs_time.png')
    plt.figure(figsize=(8, 6))
    plt.plot(data['time'], data['temperature'], label='Temperature', color='red')
    plt.xlabel('Time (s)')
    plt.ylabel('Temperature (°C)')
    plt.title('Temperature vs Time')
    plt.grid(True)
    plt.savefig(temperature_plot_path)
    plt.close()

    return voltage_plot_path, temperature_plot_path

# 生成PDF报告
def generate_pdf_report(data, plots, output_pdf):
    """
    根据实验数据生成PDF报告。
    :param data: 实验数据
    :param plots: 图表路径
    :param output_pdf: 输出的 PDF 文件路径
    """
    c = canvas.Canvas(output_pdf, pagesize=letter)
    c.setFont("Helvetica-Bold", 20)
    c.drawString(100, 750, "Battery Experiment Report")

    # 实验条件
    c.setFont("Helvetica", 12)
    c.drawString(100, 720, f"Experiment Duration: {data['time'].iloc[-1]} seconds")
    c.drawString(100, 700, f"Initial Voltage: {data['voltage'].iloc[0]} V")
    c.drawString(100, 680, f"Final Voltage: {data['voltage'].iloc[-1]} V")
    c.drawString(100, 660, f"Temperature Range: {data['temperature'].min()} - {data['temperature'].max()} °C")

    # 插入图表
    c.drawString(100, 620, "Voltage vs Time Plot:")
    c.drawImage(plots[0], 100, 400, width=400, height=200)
    c.drawString(100, 380, "Temperature vs Time Plot:")
    c.drawImage(plots[1], 100, 180, width=400, height=200)

    c.save()

# 生成Word报告
def generate_word_report(data, plots, output_docx):
    """
    生成包含实验数据和图表的Word报告。
    :param data: 实验数据
    :param plots: 图表文件路径
    :param output_docx: 输出的 Word 文件路径
    """
    doc = Document()
    doc.add_heading('Battery Experiment Report', 0)

    # 实验条件
    doc.add_heading('Experiment Conditions:', level=1)
    doc.add_paragraph(f"Experiment Duration: {data['time'].iloc[-1]} seconds")
    doc.add_paragraph(f"Initial Voltage: {data['voltage'].iloc[0]} V")
    doc.add_paragraph(f"Final Voltage: {data['voltage'].iloc[-1]} V")
    doc.add_paragraph(f"Temperature Range: {data['temperature'].min()} - {data['temperature'].max()} °C")

    # 插入图表
    doc.add_heading('Voltage vs Time Plot:', level=1)
    doc.add_picture(plots[0], width=400, height=200)
    doc.add_heading('Temperature vs Time Plot:', level=1)
    doc.add_picture(plots[1], width=400, height=200)

    doc.save(output_docx)

def generate_report(file_path, output_dir, output_pdf='battery_report.pdf', output_docx='battery_report.docx'):
    """
    自动生成电池实验报告，包括实验数据和图表。
    :param file_path: 实验数据文件路径
    :param output_dir: 输出文件夹
    :param output_pdf: 输出的 PDF 文件路径
    :param output_docx: 输出的 Word 文件路径
    """
    data = load_experiment_data(file_path)
    
    # 确保输出目录存在
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # 生成图表
    plots = generate_plots(data, output_dir)

    # 生成 PDF 和 Word 报告
    generate_pdf_report(data, plots, output_pdf)
    generate_word_report(data, plots, output_docx)

    print(f"报告已生成：PDF - {output_pdf}, Word - {output_docx}")
